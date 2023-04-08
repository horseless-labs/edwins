from docx.api import Document
import sqlite3
import os

src_file = "data/test.docx"
db_file = "../volunteer_intake_flask/instance/intake_new.sqlite"

def get_db(db_file):
    conn = sqlite3.connect(db_file)
    conn.row_factory = sqlite3.Row
    return conn

def open_document(src_file):
    try:
        doc = Document(src_file)
    except FileNotFoundError:
        print(f"The file '{src_file}' could not be found")
    except Exception as e:
        print(f"An error occurred while trying to open {src_file}: {e}")
    else:
        print(f"Opening {src_file}")
        return doc

def extract_fields(document):
    form = []
    for table in document.tables:
        for row in table.rows:
            #ueg = "NEW TABLE\n"
            ueg = "|".join([cell.text for cell in row.cells])
            if ueg.endswith("|"):
                ueg += "NA\n"
            else:
                ueg += '\n'
            form.append(ueg)
    return form

def extract_paragraphs(document):
    doc_length = len(document.paragraphs)
    paras = ""
    for p in range(doc_length):
        para = document.paragraphs[p]
        paras += para.text + '\n'
    return paras

"""
Noticing now that the extracted fields so far do not include check boxes.
TODO:
    - Add a dictionary
"""
def parse_fields(field_str):
    rows = field_str.split('\n')
    print(rows)

    text_areas = []
    for row in row:
        if "|" not in row:
            text_areas.append(row)
            continue

# My first way of approaching this takes the complete form and parses individual
# lines with underscores. Messy and unpythonic.
def parse_check_boxes(paragraphs):
    paragraphs = paragraphs.split('\n')
    boxes = []
    for para in paragraphs:
        if '_' in para and para.count('_') < 10:
            boxes.append(para)
    return boxes

# Reduces a list of check boxes to a small number of strings, one for each of the
# broader categories that they are found in on the volunteer intake form.
def stringify_checkbox(checkbox):
    str_checkbox = ', '.join(checkbox)
    return str_checkbox

def reduce_boxes(boxes):
    selected_interests = stringify_checkbox(boxes[:7])
    # Things like other_interests are not collected from checkboxes.

    oef = stringify_checkbox(boxes[7:12])
    students_lives = stringify_checkbox(boxes[12:19])
    class_education = stringify_checkbox(boxes[19:23])
    facilities = stringify_checkbox(boxes[23:32])
    clerical_advo = stringify_checkbox(boxes[32:40])
    return [selected_interests, oef, students_lives, class_education, facilities, clerical_advo]

def break_pipe(str):
    return str.split('|')[1]

def reduce_form(form):
    """
    for i in range(len(form)):
        print(f"{i}: {form[i]}")
    """
    active = 1
    name = break_pipe(form[1])
    address = break_pipe(form[2])

    location = break_pipe(form[3])
    location = location.split(', ')
    print(location)
    city = location[0]
    state = location[1]
    zip = location[2]

    home_phone = break_pipe(form[4])
    occupation = break_pipe(form[5])
    employer = break_pipe(form[6])
    cell_phone = break_pipe(form[7])
    email = break_pipe(form[8])
    dob = break_pipe(form[9])

    other_interests = form[18]
    skills = form[19]
    experience = form[20]
    return [active, name, address, city, state, zip, home_phone, occupation,
            employer, cell_phone, email, dob, other_interests, skills, experience]

# Kludgy, but something is up with the Word document.
def get_availability_info(paragraphs):
    avail_start_idx = paragraphs.find("How many hours per week are you interested in volunteering? ")
    avail_end_idx = paragraphs.find("Preferred location for volunteer assignments? ")
    avail = paragraphs[avail_start_idx:avail_end_idx].strip()
    avail = avail.split('\n')[1]
    
    location_start_idx = paragraphs.find("Preferred location for volunteer assignments? ")
    location_end_idx = paragraphs.find("Preferred time to volunteer? ")
    location = paragraphs[location_start_idx:location_end_idx]
    location = location.split('\n')[1]

    times_start_idx = paragraphs.find("Preferred time to volunteer?")
    times_end_idx = paragraphs.find("Interests")
    times = paragraphs[times_start_idx:times_end_idx]
    times = times.split('\n')[1]

    return [avail, location, times]

if __name__ == '__main__':
    doc = open_document(src_file)

    if doc:
        # Intermediate step
        paragraphs = extract_paragraphs(doc)

        # Check boxes
        boxes = parse_check_boxes(paragraphs)
        form = extract_fields(doc)

        boxes = reduce_boxes(boxes)
        form = reduce_form(form)
        
        print(len(form))
        for i in range(len(form)):
            print(f"{i}: {form[i]}")
        

        avail = get_availability_info(paragraphs)
        other_students_lives = "NA"
        guest_speaker = "NA"
        
        db = get_db(db_file)

        # Use a dummy document that has all fields filled out.
        # Stuff like other_students_lives is not represented here, and it's breaking the parser
        try:
            db.execute(
                    "INSERT INTO volunteer (active, name, address, city, state, zip, home_phone, occupation, employer, cell_phone, email, dob, \
                    availability, location, times, selected_interests, other_interests, skills, experience, oef, students_lives, other_students_lives, class_education, \
                    guest_speaker, facilities, clerical_advo) \
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
                    (form[0], form[1], form[2], form[3], form[4], form[5], form[6], form[7], form[8], form[9], form[10],
                     form[11], form[12], avail[0], avail[1], avail[2], boxes[0], form[13], form[14], boxes[1],
                     boxes[2], other_students_lives, boxes[3], guest_speaker, boxes[4], boxes[5])
            )
            db.commit()
        except db.IntegrityError:
            error = f"User {form[1]} is already registered"